#!/usr/bin/env python3
"""Generate SCREENSHOT_TEMPLATE.docx — a drop-in template for the screenshot PDF.

One page per shot: a heading, the caption, and an empty bordered box to paste the
screenshot into. Open in Word, drag a screenshot into each box, then File > Save
As / Export > PDF.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOREST = RGBColor(0x2E, 0x58, 0x44)
HONEY = RGBColor(0xC9, 0x98, 0x46)
GREY = RGBColor(0x88, 0x88, 0x88)

# (number, title, caption, group, flag)  flag: '' | 'secure' | 'cond'
SHOTS = [
    (1, "Onboarding", "First-run setup: welcome trust pills, then username, sober/quit date, daily spend + currency, and notification preferences. Multi-page — one box per step.", "Entry / first-run", "cond"),
    (2, "Lock screen", "PIN pad / fingerprint prompt shown when the app opens locked.", "Entry / first-run", "cond"),
    (3, "Welcome dialog", "“A note from Shawn” — one-time popup on first Home visit.", "Entry / first-run", "cond"),
    (4, "Safety disclaimer dialog", "One-time medical/crisis disclaimer with a “Get crisis help” button.", "Entry / first-run", "cond"),

    (5, "Home — Serenity hero card", "The centrepiece: a plant that grows with your sobriety stage, inside a decorative arch, with a live Days / Hours / Minutes / Seconds counter.", "1 · Home tab", ""),
    (6, "Home — Money saved card", "Running total of money not spent. Only appears if a daily spend is set.", "1 · Home tab", "cond"),
    (7, "Home — Journey card", "Milestone path (Day 0 → 7 → 30 → 90 → 180 → 365). Taps through to the Milestone screen.", "1 · Home tab", ""),
    (8, "Home — Daily pledge card", "Type today’s pledge; shows your pledge streak.", "1 · Home tab", ""),
    (9, "Home — Intention card", "Set today’s intention / check-in.", "1 · Home tab", ""),
    (10, "Home — Gratitude card", "Add a gratitude note for today.", "1 · Home tab", ""),
    (11, "Home — My Reason card", "Your personal reason for staying sober.", "1 · Home tab", ""),
    (12, "Home — Today’s Session card", "Today’s planned workout from the Planner. Only appears if a training plan is active.", "1 · Home tab", "cond"),
    (13, "Home — Weekly Goals card", "Checklist of weekly goals. Only appears if goals are set.", "1 · Home tab", "cond"),
    (14, "Home — Daily Missions card", "Three small daily missions to tick off.", "1 · Home tab", ""),
    (15, "Home — Check-in card", "Four quick-log buttons (Craving, Thought, Activity, Sleep), each opens a bottom sheet.", "1 · Home tab", ""),
    (16, "Sheet — Log a Craving", "Intensity, trigger, notes.", "1 · Home tab", ""),
    (17, "Sheet — Log a Thought", "Quick CBT-style reframe.", "1 · Home tab", ""),
    (18, "Sheet — Log an Activity", "Record a movement / activity.", "1 · Home tab", ""),
    (19, "Sheet — Log Sleep", "Record last night’s sleep.", "1 · Home tab", ""),
    (20, "Home — Today’s Reminder", "The daily quote card.", "1 · Home tab", ""),
    (21, "Home — Recovery banner", "Body-healing progress. Taps through to the Recovery timeline.", "1 · Home tab", ""),

    (22, "Progress screen", "Milestones, insight cards (Craving/Sleep/Movement/Thoughts), What I’ve Learned, Tender Hours, cravings heatmap, Recovery Capital.", "2 · Progress tab", ""),
    (23, "Insights screen", "Deeper charts / analytics of your logged data.", "2 · Progress tab", ""),
    (24, "Heatmap / Recovery Map", "Calendar heatmap of activity; tap a day for its detail.", "2 · Progress tab", ""),
    (25, "History screen", "All logged entries with filter tabs: Cravings, Thoughts, Activity, Sleep, Journal, Gratitude, All.", "2 · Progress tab", ""),
    (26, "Milestone screen", "The milestone celebration / shareable milestone card.", "2 · Progress tab", ""),
    (27, "Weekly Care Summary", "A week-in-review summary card (shareable).", "2 · Progress tab", ""),
    (28, "Recovery timeline", "Body-healing milestones from 12 hours → 1 year+ (M1–M11), each with what’s happening in your body.", "2 · Progress tab", ""),

    (29, "Toolkit home", "The 3× grid of 11 coping tools + (if set) a green Call-contact button.", "3 · Toolkit tab", ""),
    (30, "Breathing — choose", "Recommended pattern card + a 2×2 grid (Box, Night, Coherent, Ocean).", "3 · Toolkit tab", ""),
    (31, "Breathing — all patterns", "Library of 15 patterns (Box, 4-7-8, Calm, Power, Reset, Triangle, Anchor, Rescue, Ocean, Morning, Coherent, 6-2-8, Square+, Warrior, Night).", "3 · Toolkit tab", ""),
    (32, "Breathing — session", "The animated expanding/contracting breath ring with phase pills and timer.", "3 · Toolkit tab", ""),
    (33, "Meditation", "Guided list: Urge Surfing (live, with audio player), Body Scan, Gratitude Reset, Safe Place, Self-Compassion (coming soon).", "3 · Toolkit tab", ""),
    (34, "CBT guides", "Five quick reframe guides (capture the list + one opened guide).", "3 · Toolkit tab", ""),
    (35, "My Reasons", "Your “why I’m doing this” list.", "3 · Toolkit tab", ""),
    (36, "HALT", "Hungry / Angry / Lonely / Tired self-check with advice for each.", "3 · Toolkit tab", ""),
    (37, "Play the Tape", "“What happens if I drink” vs “if I stay sober” + what-helps.", "3 · Toolkit tab", ""),
    (38, "Mindfulness", "Six grounding exercises (5-senses, breath, body, etc.).", "3 · Toolkit tab", ""),
    (39, "Urge Timer", "Ride-out-the-urge countdown timer.", "3 · Toolkit tab", ""),
    (40, "TIPP", "The DBT TIPP crisis skill (Temperature, Intense exercise, Paced breathing, Paired muscle relaxation).", "3 · Toolkit tab", ""),
    (41, "Puzzle", "A distraction puzzle/game + grounding “right now” steps.", "3 · Toolkit tab", ""),
    (42, "100-Day Challenge", "A 100-square grid you fill in daily, with stickers and a shareable card.", "3 · Toolkit tab", ""),

    (43, "Journal list", "Your entries with a “+” to add (Plain entry vs Daily Reflection).", "4 · Journal tab (screenshot-protected)", "secure"),
    (44, "New entry — mood selector", "Mood scale: Great · Good · Okay · Hard · Crisis.", "4 · Journal tab (screenshot-protected)", "secure"),
    (45, "New entry — prompt packs", "Choose a prompt set: Reflection · Gratitude · Hard day · Wins · Craving · People (with voice-to-text).", "4 · Journal tab (screenshot-protected)", "secure"),
    (46, "Daily Reflection template", "Guided sections: Mood, Grateful, Anchors, Wins, Cravings, Intention, Affirmation.", "4 · Journal tab (screenshot-protected)", "secure"),
    (47, "Journal entry detail", "Viewing / editing a saved entry.", "4 · Journal tab (screenshot-protected)", "secure"),
    (48, "Vision Board", "Grid of “dreams” (text + gallery photos); add/edit a dream; suggested affirmations.", "4 · Journal tab (screenshot-protected)", "secure"),
    (49, "Vision detail", "A single vision item full-screen.", "4 · Journal tab (screenshot-protected)", "secure"),
    (50, "Affirmations / Zen exercises", "Custom affirmations and short zen exercises in this tab.", "4 · Journal tab (screenshot-protected)", "secure"),

    (51, "Planner ▸ Overview", "Today’s session, active goal cards with countdowns, planned-vs-actual.", "5 · Planner tab", ""),
    (52, "Planner ▸ Calendar", "Month grid. Markers: flag (goal/event day), dot (upcoming session), checkmark (completed), top-right dot (multiple).", "5 · Planner tab", ""),
    (53, "Planner ▸ Streaks", "Training streaks / consistency.", "5 · Planner tab", ""),
    (54, "Add / Edit Goal", "Goal type Exercise or Weight, name, target date.", "5 · Planner tab", ""),
    (55, "Add / Log Session", "Pick a session type (Easy run, Intervals, Tempo, Long run, Rest, Cross-train, Swim, Ride, Walk, Hike, Gym, Yoga, Cardio, Other).", "5 · Planner tab", ""),
    (56, "Log Activity", "Record a completed activity.", "5 · Planner tab", ""),
    (57, "Body Journey", "Weight / body tracking over time (chart).", "5 · Planner tab", ""),
    (58, "Planner History", "Past sessions log.", "5 · Planner tab", ""),
    (59, "Planner Insights", "Training stats / trends.", "5 · Planner tab", ""),
    (60, "Planner Share card", "Shareable training summary image.", "5 · Planner tab", ""),

    (61, "Profile / Settings", "Grouped rows: Profile (username, sober date, reasons, weekly goals), Money (spend, currency, savings goal), Reminders (notifications, schedule, diagnostics), Preferences (haptics, units, appearance, language), Safety (emergency contact), and links.", "6 · Profile tab", ""),
    (62, "Backup & Restore", "Export (optional passphrase encryption), Restore, and “What’s included”.", "6 · Profile tab", ""),
    (63, "Privacy", "The in-app privacy policy.", "6 · Profile tab", ""),
    (64, "Support Groups", "Directory of recovery fellowships (AA, NA, SMART Recovery, etc.) with descriptions.", "6 · Profile tab", ""),
    (65, "Meetings", "Your meeting reminders (title, time, location); add / empty states.", "6 · Profile tab", ""),
    (66, "Crisis help", "Crisis lines / emergency signpost.", "6 · Profile tab", ""),
    (67, "Future Letter", "Write a letter to your future self, delivered on a chosen day.", "6 · Profile tab", ""),
    (68, "Pre-Craving Plan", "Build an “if a craving hits” action plan + a guided runner.", "6 · Profile tab", ""),
    (69, "What I’ve Learned", "Your saved lessons / insights, shareable.", "6 · Profile tab", ""),
    (70, "Thought Record", "A 4-step CBT thought-record flow.", "6 · Profile tab", ""),
    (71, "About", "Version, credits, links.", "6 · Profile tab", ""),
]


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def box_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'dashed')
        e.set(qn('w:sz'), '12')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'BBBBBB')
        borders.append(e)
    tcPr.append(borders)


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.6)
        s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # ── Title page ───────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Journey Forward')
    r.font.size = Pt(34)
    r.font.bold = True
    r.font.color.rgb = FOREST
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Screen-by-Screen Walkthrough')
    r.font.size = Pt(16)
    r.font.color.rgb = HONEY
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run(
        'Drop one screenshot into each box, then export to PDF '
        '(File › Save As / Export ‣ PDF).').font.color.rgb = GREY
    warn = doc.add_paragraph()
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wr = warn.add_run(
        '⚠  The Journal tab (Shots 43–50) is screenshot-protected by '
        'Android — capture those by photographing the screen with another '
        'device.')
    wr.font.size = Pt(9)
    wr.italic = True
    wr.font.color.rgb = RGBColor(0xA0, 0x6A, 0x00)
    doc.add_page_break()

    current_group = None
    for num, title, caption, group, flag in SHOTS:
        if group != current_group:
            current_group = group
            h = doc.add_paragraph()
            hr = h.add_run(group.upper())
            hr.font.size = Pt(13)
            hr.font.bold = True
            hr.font.color.rgb = FOREST
            # underline rule
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'C99846')
            pbdr.append(bottom)
            pPr.append(pbdr)

        # Shot heading
        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(8)
        hr = head.add_run('Shot %d — %s' % (num, title))
        hr.font.size = Pt(14)
        hr.font.bold = True
        hr.font.color.rgb = FOREST
        if flag == 'secure':
            tag = head.add_run('   ⚠ screenshot-protected')
            tag.font.size = Pt(9)
            tag.italic = True
            tag.font.color.rgb = RGBColor(0xA0, 0x6A, 0x00)
        elif flag == 'cond':
            tag = head.add_run('   (only in some states)')
            tag.font.size = Pt(9)
            tag.italic = True
            tag.font.color.rgb = GREY

        # Caption
        cap = doc.add_paragraph()
        cr = cap.add_run(caption)
        cr.font.size = Pt(10.5)
        cr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # Empty placeholder box (1x1 table)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        cell.width = Cm(8.5)
        tbl.rows[0].height = Cm(11.0)
        from docx.enum.table import WD_ALIGN_VERTICAL
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        box_borders(cell)
        shade(cell, 'FAF8F5')
        ph = cell.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr = ph.add_run('▢  paste screenshot here')
        pr.font.size = Pt(11)
        pr.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

        doc.add_page_break()

    out = 'SCREENSHOT_TEMPLATE.docx'
    doc.save(out)
    print('wrote %s  (%d shot pages)' % (out, len(SHOTS)))


if __name__ == '__main__':
    main()
